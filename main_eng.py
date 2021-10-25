import docx
from regex_function import *
import pandas as pd
from nltk.tokenize import sent_tokenize
import pdftotext
import re


# doc = docx.Document("op20190216US-EN.docx")

# full_list = []
# for i in doc.paragraphs:
#     full_list.append(i.text)
# with open("OP20200144EP_영어번역문.docx", "rb") as f:
#     pdf = pdftotext.PDF(f)
#
# pdfToken_eng = []
# for page in pdf:
#     pdfToken_eng.extend(page.strip().lstrip().split('\n'))

doc = docx.Document("20191022_ox20190011kr_EN-rev.docx")

full_list = []
for i in doc.paragraphs:
    full_list.append(i.text.strip())
# print(full_list)
# print(pdfToken_eng)
# 4. 마지막 인덱스 구하기
def find_last_index(whole_list, regex):
    last_idx = -1
    for idx in range(len(whole_list)):
        check = re.compile(regex)
        check_word = check.search(whole_list[idx])
        if bool(check_word):
            last_idx = idx
            continue
    return last_idx


# 4-1. 메인 로직
def header_point_idx(whole_list, last_idx):
    modified_result = []
    for idx in range(len(whole_list)):
        if idx <= last_idx:  # first_idx = 1
            modified_result.append('')
        else:
            modified_result.append(whole_list[idx])
    return modified_result


def footer_point_idx(whole_list, last_idx):
    modified_result = []
    for idx in range(len(whole_list)):
        if idx >= last_idx:  # first_idx = 1
            modified_result.append('')
        else:
            modified_result.append(whole_list[idx])
    return modified_result


def del_header(keywords, regex):
    last_idx = find_last_index(keywords, regex)
    modified_result = header_point_idx(keywords, last_idx)
    return modified_result  # header 날린 값 / last_idx


def del_footer(keywords, regex):
    last_idx = find_last_index(keywords, regex)
    modified_result = footer_point_idx(keywords, last_idx)
    return modified_result


# 4-2. 지울 header, footer 담기
pure_main_footer = del_header(full_list, "Technical Field")
mod_header_footer_eng = del_footer(pure_main_footer, "CLAIMS")

# print(mod_header_footer_eng)


# # 5. 규칙에 맞는 정규식 함수 돌리기



modified_pdf_eng = []  # 함수돌고나온거
for sentence in mod_header_footer_eng:
    # sentence = delete_sub(sentence)
    sentence = delete_eng_small_header(sentence.strip())

    sentence = delete_eng_page_num(sentence)

    sentence = delete_eng_only_num(sentence)

    sentence = delete_page_footer(sentence)

    # sentence = delete_bracket(sentence)

    modified_pdf_eng.append(sentence.strip())

# print(modified_pdf_eng)
whole_sentence_eng = " ".join(modified_pdf_eng)

split_sentence_eng = re.split(r'[[\d]+[]]', whole_sentence_eng.strip())

result_list_eng = []
for sentence in split_sentence_eng:

    sentence = delete_big_bracket(sentence)

    sentence = delete_multi_space(sentence)

    result_list_eng.append(sentence)

result_list_eng = [v for v in result_list_eng if v]


# bracket_list_eng = []
# for sentence in result_list_eng:
#
#     sentence = delete_sub_eng(sentence)
#     bracket = append_bracket(sentence)
#     bracket_list_eng.append(bracket)


# add_dash_result_eng = []
#
# for sentence in result_list_eng:
#
#     sent_tokens_list_eng = para_nltk_sentence(sentence)
#
#     for sent_tokens in sent_tokens_list:
#         sent_tokens = delete_sub_eng(sent_tokens)
#         add_dash_result_eng.append(sent_tokens)
#
#
# bracket_list = []
# add_dash_result_eng = [v for v in add_dash_result_eng if v]
# # print(add_dash_result_eng)
# for i, sentence in enumerate(add_dash_result_eng):
#     no_list_eng.append(i + 1)
#
#     sentence, bracket = append_bracket(sentence)
#     bracket_list.append(bracket)
#
#
# com_list_eng = []
# for com in bracket_list:
#     com_list_eng.append(str("".join(com)))

# print(com_list_eng)
#add_dash_result_eng 영문

# print(add_dash_result_eng)
# print(len(add_dash_result_eng))
# print(len(com_list_eng))
# print(len(result_list_eng))
# print(len(no_list))
# print(len(result_list))

# print(result_list_eng)
# print(len(no_list))
# print(len(result_list_eng))
# print(len(com_list))
# # 7. 코멘트 함수 풀기

# 8. 엑셀 만들기
# df = pd.DataFrame({'NO': no_list, '원문': result_list_eng, '병기문자': com_list_eng})
# df_i = df.set_index('NO', inplace=True)
#
# df.to_excel('영문테스트-EN.pdf.xlsx')
