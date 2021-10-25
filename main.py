import docx
from regex_function import *
from itertools import zip_longest


from main_eng import *
import pandas as pd
from nltk.tokenize import sent_tokenize
import pdftotext
import re

# TODO list 와 string


# doc = docx.Document("op20190216US-EN.docx")

# full_list = []
# for i in doc.paragraphs:
#     full_list.append(i.text)
# with open("OP20200144EP_국문 명세서.docx", "rb") as f:
#     pdf = pdftotext.PDF(f)
#
# pdfToken = []
# for page in pdf:
#     pdfToken.extend(page.strip().lstrip().split('\n'))
#
#
doc = docx.Document("OP20200144EP_국문 명세서.docx")

full_list = []
for i in doc.paragraphs:
    full_list.append(i.text.strip())

# print(pdfToken)
# 4. 마지막 인덱스 구하기
def find_last_index(whole_list, regex):
    last_idx = -10
    for idx in range(len(whole_list)):
        check = re.compile(regex)
        check_word = check.search(whole_list[idx])
        if bool(check_word):
            last_idx = idx
            continue
    return last_idx


# 4-1. 메인 로직
def header_point_idx(whole_list, last_idx):
    modified_result = []
    for idx in range(len(whole_list)):
        if idx <= last_idx:  # first_idx = 1
            modified_result.append('')
        else:
            modified_result.append(whole_list[idx])
    return modified_result


def footer_point_idx(whole_list, last_idx):
    modified_result = []
    for idx in range(len(whole_list)):
        if idx >= last_idx:  # first_idx = 1
            modified_result.append('')
        else:
            modified_result.append(whole_list[idx])
    return modified_result


def del_header(keywords, regex):
    last_idx = find_last_index(keywords, regex)
    modified_result = header_point_idx(keywords, last_idx)
    return modified_result  # header 날린 값 / last_idx


def del_footer(keywords, regex):
    last_idx = find_last_index(keywords, regex)
    modified_result = footer_point_idx(keywords, last_idx)
    return modified_result


# 4-2. 지울 header, footer 담기
pure_main_footer = del_header(full_list, "【기술분야】")
mod_header_footer = del_footer(pure_main_footer, "【청구범위】")


# # 5. 규칙에 맞는 정규식 함수 돌리기
# print(mod_header_footer)


modified_pdf = []  # 함수돌고나온거
# for sentence in mod_header_footer:
#
#     # sentence = delete_sub(sentence)
#     #
#     # sentence = delete_page_num(sentence)
#     #
#     # sentence = delete_page_footer(sentence)
#
#     # sentence = delete_bracket(sentence)
#
#     modified_pdf.append(sentence.strip())

# 【0001】
whole_sentence = " ".join(modified_pdf)

split_sentence = re.split(r'([【][\d]+[】]+)', whole_sentence.strip())
print(split_sentence)
result_list = []

for sentence in split_sentence:

    sentence = delete_big_bracket(sentence)

    sentence = delete_multi_space(sentence)

    result_list.append(sentence)


result_list = [v for v in result_list if v]
# print(result_list)
add_dash_result = [] # add_dash_result_ko와eng 묶은 zip
add_dash_result_ko = []
add_dash_result_eng = []

for ko_par in result_list:
    add_dash_result_ko.append('--------------------')
    add_dash_result_ko.append(ko_par)
    add_dash_result_ko.append('--------------------')
    #
    # # add_dash_result_ko.append(kor_list)
    # # add_dash_result_eng.append(eng_list)
    # for ko in kor_list:
    #     add_dash_result_ko.append(ko)
    # for en in eng_list:
    #     add_dash_result_eng.append(en)
    #     add_dash_result.append(list(zip_longest(kor_list, eng_list, fillvalue='-')))
    #     add_dash_result.append('--------------------')
# print(add_dash_result_ko)
# print(add_dash_result_ko)
comment_ko = []
comment_en = []
no_list = []
# for ko in add_dash_result_ko:
#     # no_list.append(i+1)
#     bracket_ko = find_bracket(str(ko))
#
#     comment_ko.append(bracket_ko)
#
#
# for en in add_dash_result_eng:
#     # no_list.append(i+1)
#
#     bracket_en = find_bracket(str(en))
#
#     comment_en.append(bracket_en)
#
#
# comment_list_kor=[]
# comment_list_eng=[]
# comment_list=[]
# for z in add_dash_result:
#     for i in z:
#             for k in i:
#                 print(k)
#
#     zip = zip_longest(no_list, add_dash_result, comment, fillvalue='-')
#
#
#   for ko, en in kor_list, eng_list:
#         bracket_ko = append_bracket(ko)
#         bracket_eng = append_bracket(en)
#         comment_list_kor.append(bracket_ko)
#         comment_list_eng.append(bracket_eng)
#         sentence, bracket = append_bracket(sentence)
#         comment_list.append(bracket)
#
#
# bracket_list = []
#
# for i, sentence in enumerate(add_dash_result):
#     no_list.append(i + 1)
#
#     sentence, bracket = append_bracket(sentence)
#
#     bracket_list.append(bracket)
#
# com_list_ko = []
# for com in comment_ko:
#     com_list_ko.append(str("".join(com)))
#
# com_list_en = []
# for com in comment_en:
#     com_list_en.append(str("".join(com)))
#
# # # print(com_list)
# # #
# # # print(len(add_dash_result))
# # #
# # # print(len(no_list))
# # # # print(len(result_list))
# # # print(len(com_list))
# # # print(len(result_list_eng))
# # # print(len(com_list_eng))
# #
#
# for i in range(len(add_dash_result_eng)):
#     no_list.append(i+1)
#
# zipped = zip_longest(no_list, add_dash_result_ko, com_list_ko, add_dash_result_eng, com_list_en, fillvalue='-')
#
# a = []
# b = []
# c = []
# d = []
# e = []
# for aa, bb, cc, dd, ee in list(zipped):
#     a.append(aa)
#     b.append(bb)
#     c.append(cc)
#     d.append(dd)
#     e.append(ee)
#
# # # print(b)
# # # # 8. 엑셀 만들기
# df = pd.DataFrame({'NO':a , '국문': b, '국문병기': c, '영문': d, '영문병기': e})
# df_i = df.set_index('NO', inplace=True)
#
# df.to_excel('OP20200144EP_국문 명세서.docx.xlsx')
